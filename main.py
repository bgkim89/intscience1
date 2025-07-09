import streamlit as st
from docx import Document
import pandas as pd
import re
from io import BytesIO

def extract_info_from_word(doc_file):
    """
    Word 문서에서 페이지별 정보를 추출합니다.
    """
    document = Document(doc_file)
    extracted_data = []

    # 각 단락을 순회하며 페이지 분리 및 정보 추출 시도
    current_page_number = 0
    current_page_data = {"five_digit_number": None, "table_row1_col1": None, "table_row2_col1": None}
    
    # Simple heuristic to identify "pages" based on section breaks or similar logic
    # This is a simplification. A robust page detection in DOCX is complex and might require
    # rendering or more advanced parsing. Here, we assume a new 'page' often starts with
    # a new table or a significant break after a table.

    for i, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()

        # Try to find a five-digit number (e.g., invoice number, document ID)
        # We assume this number appears *before* the table for that page.
        five_digit_match = re.search(r'\b(\d{5})\b', text)
        if five_digit_match and not current_page_data["five_digit_number"]:
            current_page_data["five_digit_number"] = five_digit_match.group(1)

        # Iterate through tables to find the relevant data
        # This part is tricky as tables are not directly linked to paragraphs in terms of 'page'
        # A more robust solution might involve analyzing layout or using external libraries
        # For simplicity, we'll process tables as they appear and associate them with the
        # most recently found five-digit number.
        for table in document.tables:
            # Check if this table has at least 2 rows and 1 column
            if len(table.rows) >= 2 and len(table.rows[0].cells) >= 1 and len(table.rows[1].cells) >= 1:
                try:
                    table_data_row1_col1 = table.rows[0].cells[0].text.strip()
                    table_data_row2_col1 = table.rows[1].cells[0].text.strip()
                    
                    # If we found a table and potentially a five-digit number for this 'page'
                    if table_data_row1_col1 and table_data_row2_col1 and current_page_data["five_digit_number"]:
                        # Append the collected data and reset for the next 'page'
                        extracted_data.append({
                            "five_digit_number": current_page_data["five_digit_number"],
                            "table_row1_col1": table_data_row1_col1,
                            "table_row2_col1": table_data_row2_col1
                        })
                        # Reset for the next potential 'page' or set of data
                        current_page_data = {"five_digit_number": None, "table_row1_col1": None, "table_row2_col1": None}
                        # We break here assuming one relevant table per "page" for this logic
                        break 
                except IndexError:
                    # Handle cases where cell might not exist (though checked with len)
                    pass
        
    return extracted_data

st.set_page_config(page_title="Word to CSV Converter", layout="centered")

st.title("📄 Word 파일을 CSV로 변환하기")

st.write(
    """
    이 앱은 Word 문서(`.docx`)를 업로드하면 각 페이지의 특정 내용을 추출하여 CSV 파일로 변환합니다.
    """
)

st.warning(
    """
    **중요:** 이 프로그램은 Word 문서의 "페이지" 개념을 내부적으로 정확히 파악하기 어렵습니다. 
    대신, **표 위에 있는 다섯 자리 숫자**를 기준으로 새로운 "페이지" 데이터 블록을 식별하려고 시도합니다. 
    따라서, 각 페이지마다 표가 있고 그 위에 다섯 자리 숫자가 명확히 존재해야 정확하게 동작합니다.
    """
)

uploaded_file = st.file_uploader("Word 문서(.docx)를 업로드해주세요.", type=["docx"])

if uploaded_file:
    st.info("파일을 읽는 중입니다. 잠시만 기다려 주세요...")
    
    try:
        # docx 파일을 BytesIO 객체로 읽기
        doc_file = BytesIO(uploaded_file.getvalue())
        extracted_info = extract_info_from_word(doc_file)

        if extracted_info:
            df = pd.DataFrame(extracted_info)
            df.columns = ["A열", "B열", "C열"] # 열 이름 설정
            
            st.subheader("변환 결과 미리보기")
            st.dataframe(df)

            csv_buffer = BytesIO()
            df.to_csv(csv_buffer, index=False, encoding='utf-8-sig') # 한글 깨짐 방지
            csv_buffer.seek(0)
            
            st.download_button(
                label="CSV 파일 다운로드",
                data=csv_buffer,
                file_name="converted_data.csv",
                mime="text/csv",
            )
            st.success("Word 문서가 성공적으로 CSV로 변환되었습니다!")
            st.balloons()
        else:
            st.warning("Word 문서에서 추출할 데이터를 찾지 못했습니다. 문서 형식이 예상과 다를 수 있습니다.")

    except Exception as e:
        st.error(f"파일을 처리하는 중 오류가 발생했습니다: {e}")
        st.error("Word 문서가 손상되었거나 지원되지 않는 형식일 수 있습니다. 또는 내부 로직이 문서 구조를 파악하지 못할 수 있습니다.")

st.markdown("---")
st.write("궁금한 점이 있으시면 언제든지 문의해주세요.")
